#!/usr/bin/env python3
"""Integrate verified sample-level FAPROTAX LOSO results into submission DOCX files."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--package-dir", required=True, type=Path)
    return parser.parse_args()


def insert_paragraph_after(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = paragraph._parent.add_paragraph()
    inserted._p.getparent().remove(inserted._p)
    new_p.addnext(inserted._p)
    inserted._p.getparent().remove(new_p)
    if style:
        inserted.style = style
    if text:
        inserted.add_run(text)
    return inserted


def find_paragraph(document: Document, exact_text: str):
    matches = [p for p in document.paragraphs if p.text.strip() == exact_text]
    if len(matches) != 1:
        raise ValueError(f"Expected one paragraph matching {exact_text!r}, found {len(matches)}")
    return matches[0]


def replace_once(paragraph, old: str, new: str) -> None:
    text = paragraph.text
    if text.count(old) != 1:
        raise ValueError(
            f"Expected one occurrence of {old!r} in paragraph, found {text.count(old)}"
        )
    paragraph.text = text.replace(old, new)


def add_alt_text(inline_shape, title: str, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def update_main_manuscript(package_dir: Path) -> Path:
    path = package_dir / "01_Main_Manuscript" / "Manuscript_npj_Biofilms_Microbiomes.docx"
    document = Document(path)

    abstract = find_paragraph(document, "Abstract")._p.getnext()
    abstract_paragraph = next(p for p in document.paragraphs if p._p is abstract)
    replace_once(
        abstract_paragraph,
        "A conserved bacterial core transferred across studies, whereas the fungal core and pH-response signatures were context dependent.",
        "A conserved bacterial core transferred across studies, and 12 FAPROTAX-predicted functional labels passed strict seven-cohort leave-one-study-out validation, whereas the fungal core and pH-response signatures were context dependent.",
    )

    results_anchor = find_paragraph(
        document,
        "Full-data external validation in field rice confirms the property separation",
    )
    results_heading = insert_paragraph_after(
        results_anchor._p.getprevious().getparent()
        if False
        else find_paragraph(
            document,
            "Major sensitivity analyses confirmed that bacterial core transportability was not an artefact of feature selection or sequencing depth. A cross-study conserved bacterial backbone of 40 features (14 named genera and 26 genus-unclassified lineages) was validated across all leave-one-study-out core definitions (Supplementary Data 1 (Table S8)). Validation exceeded simple random expectations in every held-out cohort and exceeded the more stringent prevalence- and abundance-matched null in five of seven cohorts; the two non-significant cohorts are reported explicitly above (Supplementary Data 1 (Table S6)). After rarefying both bacterial and fungal tables to 5,000 reads per sample, bacteria still retained a substantially broader core than fungi (mean core size, 32.9 versus 2.5 features; Supplementary Data 1 (Table S7)), while maintaining high held-out validation rates. Thus, the cross-study conserved bacterial backbone shows broad prevalence transportability, whereas the fungal core remains narrow and limited to a small set of shared features.",
        ),
        "Sample-level functional inference identifies a cross-study stable predicted layer",
        "Heading 2",
    )
    results_text = (
        "Official FAPROTAX v1.2.12 annotation of the seven-cohort bacterial genus "
        "matrix yielded 44 represented functional groups across 265 samples. In "
        "strict leave-one-study-out validation, a function was discovered using "
        "only the six training cohorts when it occurred in at least 50% of samples "
        "within every training cohort, and it was then tested at the same threshold "
        "in the untouched held-out cohort. Across the seven folds, 84 of 90 "
        "candidate-function predictions replicated (93.3%), and 12 FAPROTAX labels "
        "qualified and validated in all seven held-out definitions (Supplementary "
        "Figure S7; Supplementary Data 2). The mean fold validation rate was 94.4%, "
        "compared with 91.0% under 999 training-prevalence-matched random draws "
        "(global empirical P = 0.001). The stable labels included broad or nested "
        "categories related to chemoheterotrophy, cellulolysis, nitrogen fixation, "
        "nitrate reduction, photoheterotrophy/phototrophy, nitrification/ammonia "
        "oxidation, fermentation, iron respiration, and predatory or exoparasitic "
        "lifestyles. Because these labels are taxonomically inferred and only a "
        "mean 27.2% of original 16S reads received at least one FAPROTAX assignment, "
        "they indicate cross-study stable predicted ecological potential rather "
        "than measured genes, transcripts, metabolites, or activity."
    )
    insert_paragraph_after(results_heading, results_text, "Normal")

    discussion_anchor = find_paragraph(
        document,
        "The fungal result requires greater restraint. Penicillium, Mortierella, Trichoderma, and Exophiala were widespread under the primary UNITE-based definition, but only Penicillium survived strict kingdom filtering. None should be equated automatically with ericoid mycorrhizal function, and genus-level ITS annotation cannot resolve strain-specific lifestyles. The primary fungal core is best treated as a candidate set, with Penicillium receiving the strongest cross-study support. Depth-controlled resampling confirmed that this narrow fungal core is not a sequencing-depth artefact (Supplementary Data 1 (Table S7)): after rarefaction to a common depth the fungal core remained far narrower than the bacterial core, so its limited breadth is not merely a depth artefact, although whether it reflects ecology or genus-level annotation instability remains to be resolved at higher taxonomic ranks. Together, higher fungal ecological turnover, stronger host- or environment-dependent filtering and ITS methodological heterogeneity are plausible contributors; the current data do not identify their relative contribution.",
    )
    discussion_text = (
        "The sample-level FAPROTAX analysis adds a cautious functional interpretation "
        "to the transferable bacterial membership pattern. Twelve database labels "
        "survived discovery in every six-cohort training set and validation in each "
        "untouched cohort, supporting recurrence of predicted functional potential "
        "rather than a function assigned from the final pooled core alone. This set "
        "is partly hierarchical: aerobic chemoheterotrophy is nested within "
        "chemoheterotrophy, aerobic ammonia oxidation within nitrification, and "
        "photoheterotrophy within phototrophy. The 12 labels therefore span "
        "approximately nine nonredundant functional families and should not be "
        "counted as independent pathways. The low and cohort-variable annotation "
        "coverage, particularly in PRJNA577971/PRJNA578171, further argues for "
        "treating the result as hypothesis-generating support for shotgun "
        "metagenomics and biochemical validation rather than as evidence of "
        "metabolic activity."
    )
    insert_paragraph_after(discussion_anchor, discussion_text, "Normal")

    limitations = next(
        p for p in document.paragraphs if p.text.startswith("The synthesis has nine principal limitations.")
    )
    replace_once(
        limitations,
        "The synthesis has nine principal limitations.",
        "The synthesis has ten principal limitations.",
    )
    replace_once(
        limitations,
        "Ninth, amplicon data do not directly establish metabolic function, interaction, or benefit to blueberry growth.",
        "Ninth, amplicon data do not directly establish metabolic function, interaction, or benefit to blueberry growth. Tenth, FAPROTAX is a taxonomy-based inference whose assignments covered a mean 27.2% of original bacterial reads and contain nested labels, so the LOSO-stable functions represent predicted ecological potential rather than independent pathways or measured activity.",
    )

    methods_anchor = find_paragraph(
        document,
        "Core microbiome and leave-one-study-out validation",
    )
    methods_body = methods_anchor._p.getnext()
    methods_body_paragraph = next(p for p in document.paragraphs if p._p is methods_body)
    methods_heading = insert_paragraph_after(
        methods_body_paragraph,
        "FAPROTAX functional inference and leave-one-study-out validation",
        "Heading 2",
    )
    methods_text = (
        "Sample-level ecological functions were inferred from the harmonized "
        "bacterial genus count matrix using the unmodified FAPROTAX v1.2.12 "
        "database. The official collapse_table.py utility was run on full taxonomic "
        "paths, retaining represented functional groups as raw assigned counts and "
        "as fractions of each original 16S library. Per-sample annotation coverage "
        "was calculated as one minus the fraction assigned to the FAPROTAX "
        "unassigned-taxa group. Functional stability was evaluated by seven-fold "
        "leave-one-study-out validation. In each fold, a function was defined using "
        "only the six training studies and qualified when its prevalence was at "
        "least 50% within every training study; it validated when prevalence was at "
        "least 50% in the untouched held-out study. The final stable set contained "
        "only functions that qualified and validated in all seven folds. For each "
        "fold, the observed validation rate was compared with 999 random function "
        "sets matched to the training-prevalence deciles of the discovered "
        "functions (seed 20260728). A global empirical P value was calculated from "
        "the macro-average validation rate across folds. Database, input-table, and "
        "metadata SHA-256 checksums, assignment reports, sample-level abundance "
        "matrices, fold-level results, and the complete null draws are provided in "
        "Supplementary Data 2 and the archived analysis output."
    )
    insert_paragraph_after(methods_heading, methods_text, "Normal")

    conclusion = next(
        p for p in document.paragraphs if p.text.startswith("This study provides a study-aware")
    )
    replace_once(
        conclusion,
        "For Vaccinium and acidic-soil microbiome research, the 40-feature bacterial core provides a defensible and empirically validated starting point for strain-resolved metagenomics, cultivation, and synthetic-community experiments - but it is not a validated bioinoculant prescription.",
        "For Vaccinium and acidic-soil microbiome research, the 40-feature bacterial core and the 12 cross-study stable FAPROTAX-predicted labels provide a defensible starting point for strain-resolved metagenomics, cultivation, biochemical assays, and synthetic-community experiments - but they are neither direct measurements of metabolic activity nor a validated bioinoculant prescription.",
    )

    document.save(path)
    return path


def update_supplementary_information(package_dir: Path) -> Path:
    path = package_dir / "04_Supplementary_Information" / "Supplementary_Information.docx"
    figure_path = (
        package_dir
        / "faprotax_samplelevel_loso"
        / "Supplementary_Figure_FAPROTAX_LOSO.png"
    )
    document = Document(path)

    contents = next(p for p in document.paragraphs if p.text.startswith("Contents:"))
    contents.text = (
        "Contents: Supplementary Notes S1-S3; Supplementary Figures S1-S7; "
        "Supplementary Data 1 (Tables S1-S8); Supplementary Data 2 "
        "(sample-level FAPROTAX and LOSO validation)."
    )

    figures_heading = find_paragraph(document, "Supplementary Figures")
    note_heading = figures_heading.insert_paragraph_before(
        "Supplementary Note S3. Sample-level FAPROTAX and strict LOSO functional validation",
        style="Heading 1",
    )
    note1 = figures_heading.insert_paragraph_before(
        "The seven-cohort bacterial genus count matrix comprised 265 samples and "
        "744 full taxonomic paths. Sample-level ecological functions were inferred "
        "with the official, unmodified FAPROTAX v1.2.12 database. Forty-four "
        "functional groups were represented. Counts were retained both as assigned "
        "read counts and as fractions of the original bacterial 16S library; mean "
        "assignment coverage was 27.2% (sample range 8.8%-54.8%).",
        style="Normal",
    )
    note2 = figures_heading.insert_paragraph_before(
        "Strict seven-fold leave-one-study-out validation excluded the held-out "
        "cohort from discovery. In each fold, a candidate function was required to "
        "occur in at least 50% of samples within every one of the six training "
        "cohorts and was considered validated when it reached the same prevalence "
        "in the untouched cohort. Eighty-four of 90 fold-level candidate predictions "
        "validated (93.3%); 12 labels qualified and validated in all seven folds. "
        "The mean fold validation rate was 94.4%, exceeding the 91.0% mean from 999 "
        "training-prevalence-decile-matched random sets (global empirical P=0.001; "
        "seed 20260728).",
        style="Normal",
    )
    figures_heading.insert_paragraph_before(
        "The all-heldout labels were aerobic chemoheterotrophy, chemoheterotrophy, "
        "predatory or exoparasitic lifestyles, cellulolysis, nitrogen fixation, "
        "nitrate reduction, photoheterotrophy, phototrophy, nitrification, aerobic "
        "ammonia oxidation, fermentation, and iron respiration. Several are nested "
        "database categories, so the 12 labels span approximately nine "
        "nonredundant functional families. These outputs are taxonomy-based "
        "predictions and do not establish genes, transcripts, metabolites, "
        "biochemical rates, or causal benefit.",
        style="Normal",
    )

    s6_legend = next(
        p for p in document.paragraphs
        if p.text.startswith("Supplementary Figure S6.")
    )
    image_paragraph = insert_paragraph_after(s6_legend, "", "Normal")
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = image_paragraph.add_run().add_picture(str(figure_path), width=Inches(6.7))
    add_alt_text(
        shape,
        "Supplementary Figure S7",
        "Three-panel figure showing FAPROTAX annotation coverage, held-out prevalence "
        "of twelve stable functional labels, and observed versus matched-null LOSO "
        "validation rates across seven cohorts.",
    )
    legend = insert_paragraph_after(
        image_paragraph,
        "Supplementary Figure S7. Sample-level FAPROTAX annotation and seven-cohort "
        "leave-one-study-out functional validation. (A) Fraction of original 16S "
        "reads assigned to at least one FAPROTAX function in each cohort; points "
        "represent samples, boxes show the interquartile range, and center lines "
        "show medians. (B) Prevalence in each untouched held-out cohort for the 12 "
        "FAPROTAX labels that qualified and validated in every fold. (C) Observed "
        "fold validation rates and means from 999 training-prevalence-decile-matched "
        "random function sets. The observed global macro-average was 0.944 versus "
        "0.910 under the matched null (empirical P=0.001). Functional labels are "
        "taxonomy-based predictions and include nested categories.",
        "Figure Legend",
    )
    legend.paragraph_format.keep_with_next = True

    data2_heading = document.add_paragraph(
        "Supplementary Data 2", style="Heading 1"
    )
    document.add_paragraph(
        "Supplementary Data 2 is supplied as an Excel workbook and contains the "
        "analysis manifest, cohort-level annotation coverage, sample-level function "
        "abundances, function prevalence by cohort, fold-level LOSO summaries, "
        "function-level validation details, the all-heldout stable functions, and "
        "the matched-null summary.",
        style="Normal",
    )

    document.save(path)
    return path


def main() -> None:
    args = parse_args()
    package_dir = args.package_dir.resolve()
    manuscript = update_main_manuscript(package_dir)
    supplement = update_supplementary_information(package_dir)
    print(manuscript)
    print(supplement)


if __name__ == "__main__":
    main()
