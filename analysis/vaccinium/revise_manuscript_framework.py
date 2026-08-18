from pathlib import Path
from docx import Document


PACKAGE = Path("submission_package_FINAL_20260719_v5")
MANUSCRIPT = PACKAGE / "01_Manuscript" / "Manuscript.docx"
CAPTIONS = PACKAGE / "06_Data_Code_Archive" / "figure_source_data" / "Figure_Captions.txt"
STATUS = PACKAGE / "10_FINAL_VERSION_STATUS.md"


def set_text(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def find_paragraph(doc, prefix):
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def replace_in_prefix(doc, prefix, old, new):
    paragraph = find_paragraph(doc, prefix)
    text = paragraph.text
    if old not in text:
        raise RuntimeError(f"Text not found in paragraph {prefix!r}: {old!r}")
    set_text(paragraph, text.replace(old, new, 1))


def main():
    backup = MANUSCRIPT.with_name("Manuscript.before_framework_revision.docx")
    if not backup.exists():
        backup.write_bytes(MANUSCRIPT.read_bytes())

    doc = Document(MANUSCRIPT)

    abstract = (
        "Background: Blueberries and related Vaccinium species depend on root-associated microorganisms in acidic soils, yet evidence is fragmented across studies, marker regions, hosts and designs, and which microbiome properties actually reproduce across studies is unclear. "
        "Methods: We built a study-aware, dual-kingdom benchmark by uniformly reprocessing paired 16S and ITS amplicons from seven Vaccinium cohorts (four discovery and three external-validation cohorts; 211 primary and 265 extended paired samples) with study-specific denoising and harmonized taxonomic references, and evaluated study effects, pH and management responses, cross-kingdom concordance and core taxa with study-held-out validation. "
        "Results: Study identity explained approximately one-third of bacterial and one-quarter of fungal community variation. Core membership was the most reproducible property: a cross-study conserved bacterial backbone of 40 features (14 named genera plus 26 genus-unclassified lineages) was retained in every held-out definition and validated at 67.8–100%, remaining above prevalence- and abundance-matched random-subset expectations in five of seven cohorts and after rarefaction, whereas the fungal core reduced to two features and was filtering-sensitive. Environmental response was less transferable: pH repeatedly shaped community composition but yielded no universally transferable pH-responsive taxa; the field bacterial pH signature was suggestive and design-confounded (quadratic R² = 0.242, permutation P = 0.025). In one 36-sample study with provisional labels, the bacterial core supported exploratory within-study cross-block discrimination (balanced accuracy 0.722, 95% CI 0.611–0.833), not portable management prediction. "
        "Conclusions: Core membership, environmental response and prediction are distinct properties. Study-held-out validation supports a conserved bacterial backbone, but context-dependent pH response, fungal filtering sensitivity and unconfirmed management labels limit universal ecological or predictive claims."
    )
    set_text(find_paragraph(doc, "Background: Blueberries"), abstract)

    highlights = (
        "Uniform reprocessing of seven paired 16S–ITS Vaccinium cohorts (265 primary-QC samples; four discovery and three external-validation cohorts) provides a study-aware, dual-kingdom benchmark of which microbiome properties transfer across studies. Core membership was the most reproducible property, with a cross-study conserved bacterial backbone validated across held-out cohorts, whereas the fungal core contracted under strict filtering. pH and cross-kingdom signals were context-dependent, and management discrimination was a provisional within-study stress test rather than portable prediction."
    )
    set_text(find_paragraph(doc, "Uniform reprocessing of seven paired"), highlights)

    introduction = (
        "Here, we uniformly reprocessed four independent Vaccinium studies containing paired 16S and ITS libraries as a discovery set and further assembled three additional independent cohorts for external validation (seven cohorts and 265 paired samples in total). Three gaps motivated the design: core membership is not equivalent to a transferable core; environmental associations rarely receive independent study-held-out validation; and classifier performance is often reported without testing study transfer. We asked seven questions: (i) how large are study effects after taxonomic harmonization; (ii) are pH responses reproducible between controlled and field settings; (iii) does management reorganize both microbial kingdoms; (iv) can cross-study core features discriminate management treatment in held-out field blocks; (v) does bacterial-fungal community concordance persist after adjustment for study-specific design factors; (vi) which genera form a cross-study conserved core under prevalence, kingdom-confidence and depth controls; and (vii) do core membership, environmental response and predictive transferability behave as distinct properties? The conceptual framework is summarized in Figure S2."
    )
    set_text(find_paragraph(doc, "Here, we uniformly reprocessed four independent"), introduction)

    replace_in_prefix(
        doc,
        "Core restriction markedly improved",
        "Core restriction markedly improved cross-block discrimination for bacteria and for the combined feature set, but not for fungi.",
        "Within this exploratory 36-sample, four-block stress test, core restriction improved cross-block discrimination for bacteria and for the combined feature set, but not for fungi; these values should not be interpreted as portable classifier performance."
    )

    replace_in_prefix(
        doc,
        "To evaluate whether the inferred core signatures",
        "Leave-one-study-out validation showed strong transportability of the bacterial core:",
        "Leave-one-study-out validation showed broad prevalence transportability of the bacterial core:"
    )
    replace_in_prefix(
        doc,
        "To evaluate whether the inferred core signatures",
        "Overall, the bacterial core behaves as a transferable backbone,",
        "Overall, the cross-study conserved bacterial backbone shows broad prevalence transportability,"
    )

    replace_in_prefix(
        doc,
        "Major sensitivity analyses confirmed",
        "A stable bacterial backbone of 40 features",
        "A cross-study conserved bacterial backbone of 40 features"
    )
    replace_in_prefix(
        doc,
        "Major sensitivity analyses confirmed",
        "Thus, the bacterial core represents a broadly transferable backbone, whereas the fungal core is genuinely narrow and limited to a small set of shared features.",
        "Thus, the cross-study conserved bacterial backbone shows broad prevalence transportability, whereas the fungal core remains narrow and limited to a small set of shared features."
    )

    replace_in_prefix(
        doc,
        "The controlled experiment demonstrated that pH",
        "Because field pH is confounded with cultivar, management, spatial structure and soil chemistry, the field pH results are treated as design-confounded and the nonlinear bacterial signature as suggestive rather than confirmatory.",
        "Because field pH is confounded with cultivar, management, spatial structure and soil chemistry, the field pH results are treated as design-confounded and the nonlinear bacterial signature as suggestive rather than confirmatory. The weaker cross-study fungal result has at least three non-exclusive explanations: higher fungal beta diversity and ecological turnover, stronger host- or environment-dependent filtering of fungal assemblages, and ITS-specific methodological heterogeneity in primer or region choice, amplification and reference coverage. The present data cannot separate these mechanisms, so fungal non-transferability should not be read as analytical failure."
    )
    replace_in_prefix(
        doc,
        "The modest taxon-level validation rate",
        "Thus, the six nominally validated bacterial genera, one fungal genus, and the bacterial composite signature are useful candidates for targeted validation, but none should yet be described as a universal pH biomarker.",
        "Thus, the six nominally validated bacterial genera, one fungal genus, and the bacterial composite signature should be treated as candidates for targeted validation, not universal pH biomarkers."
    )

    replace_in_prefix(
        doc,
        "3.3. A conserved bacterial backbone",
        "3.3. A conserved bacterial backbone carries transferable management information",
        "3.3. A cross-study conserved bacterial backbone carries conditional management information"
    )
    replace_in_prefix(
        doc,
        "Conexibacter, Candidatus Solibacter",
        "linked prediction to interpretable abundance responses.",
        "linked the within-study discrimination signal to interpretable abundance responses."
    )
    replace_in_prefix(
        doc,
        "The fungal result requires greater restraint",
        "although whether it reflects ecology or genus-level annotation instability remains to be resolved at higher taxonomic ranks.",
        "although whether it reflects ecology or genus-level annotation instability remains to be resolved at higher taxonomic ranks. Together, higher fungal ecological turnover, stronger host- or environment-dependent filtering and ITS methodological heterogeneity are plausible contributors; the current data do not identify their relative contribution."
    )
    replace_in_prefix(
        doc,
        "Two cautious implications follow",
        "pH correction should be guided by crop, soil matrix and compartment rather than a universal microbiome optimum",
        "pH management should be guided by crop, soil matrix and compartment rather than a universal microbial prescription"
    )

    conclusion = (
        "Uniform reprocessing of seven paired 16S–ITS Vaccinium cohorts provides a study-aware, dual-kingdom benchmark separating three properties of root microbiomes: core membership, environmental response and predictive transferability. Study effects were large, and core membership transferred most strongly through a cross-study conserved bacterial backbone that recurred and validated across held-out cohorts and remained robust to random-subset and rarefaction controls, whereas the genus-level fungal core was narrow and filtering-sensitive. Environmental response was context-dependent: pH effects were reproducible at the community level but yielded no universal responsive taxa, and cross-kingdom concordance depended on fungal filtering. Predictive evidence was the most conditional: the bacterial core supported within-study cross-block discrimination under provisional labels, but portable management prediction was not established. These findings show that reproducible membership does not imply reproducible response or prediction, and that transferability in plant root microbiomes should be evaluated with study-held-out designs rather than asserted as universal microbial prescriptions."
    )
    conclusion_paragraph = next(
        paragraph for paragraph in doc.paragraphs
        if "this was not tested across studies" in paragraph.text
    )
    set_text(conclusion_paragraph, conclusion)

    doc.save(MANUSCRIPT)

    caption = "Figure S2. Conceptual framework for the study-aware dual-kingdom benchmark. Study-held-out validation separates three properties: cross-study conserved core membership, context-dependent environmental response and conditional within-study predictive discrimination. The framework emphasizes that reproducible membership does not imply reproducible response or portable prediction."
    cap_text = CAPTIONS.read_text(encoding="utf-8")
    if "Figure S2. Conceptual framework" not in cap_text:
        CAPTIONS.write_text(cap_text.rstrip() + "\n\n" + caption + "\n", encoding="utf-8")

    status = STATUS.read_text(encoding="utf-8")
    status = status.replace("Authoritative package: `submission_package_FINAL_20260719_v1`", "Authoritative package: `submission_package_FINAL_20260719_v5`")
    status = status.replace("This package contains the RA-format manuscript and the latest figure outputs:", "This package contains the revised manuscript, the latest figure outputs, and the conceptual framework supplement:")
    status = status.replace("- Figure 6: latest matched-null extended-validation PDF/PNG/SVG output.", "- Figure 6: latest matched-null extended-validation PDF/PNG/SVG output.\n- Supplementary Figure S2: conceptual framework separating core membership, environmental response and predictive transferability.")
    status = status.replace("The package is content-final for the current analysis.", "The package is content-final for the current analysis; wording has been tightened to distinguish a cross-study conserved bacterial backbone from a universal core, and to frame management discrimination as an exploratory within-study stress test.")
    STATUS.write_text(status, encoding="utf-8")

    print(MANUSCRIPT)
    print(CAPTIONS)
    print(STATUS)


if __name__ == "__main__":
    main()
